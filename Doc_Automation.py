import logging
import os
import time
import datetime
import traceback
import warnings
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
import pandas as pd
from docx.shared import Pt

Start_time = datetime.datetime.now()
path = os.getcwd()
print("Process Started!!")

logging.basicConfig(filename='Error_File.log', level=logging.DEBUG)
warnings.filterwarnings("ignore")


if os.path.exists(path + "\\Output") and os.path.isdir(path + "\\Output"):
    pass
else:
    os.mkdir(path + "\\Output")

today = datetime.datetime.now().strftime(" %m_%d_%Y  %H_%M_%S")
today_dt= datetime.datetime.now().strftime("%m_%d_%Y")

try:
    df_client_requirement = pd.read_excel(path + "\\Client Requirements.xlsx")
except:
    logging.exception("\n \n \n Error Logged: Client Requirement excel file not avaialable")


try:
    df_client_requirement["Draft Date"] = pd.to_datetime(df_client_requirement["Draft Date"])
    df_client_requirement["Draft Date"] = pd.to_datetime(df_client_requirement["Draft Date"]).dt.date

    df_client_requirement["Review Date"] = pd.to_datetime(df_client_requirement["Review Date"])
    df_client_requirement["Review Date"] = pd.to_datetime(df_client_requirement["Review Date"]).dt.date

    df_client_requirement["Approval Date"] = pd.to_datetime(df_client_requirement["Approval Date"])
    df_client_requirement["Approval Date"] = pd.to_datetime(df_client_requirement["Approval Date"]).dt.date

    Draft_Date = df_client_requirement["Draft Date"].iloc[0]
    Review_Date = df_client_requirement["Review Date"].iloc[0]
    Approval_Date = df_client_requirement["Approval Date"].iloc[0]

    Organization_Name = df_client_requirement["Organization Name"].iloc[0]
except:
    logging.exception("\n \n \n Error Logged: Error in Client Requirement file Data")

output =path + "\\Output\\"+ Organization_Name + " - " + today

if os.path.exists(output) and os.path.isdir(output):
    pass
else:
    os.mkdir(output)

try:
    Draft_Date = Draft_Date.strftime("%d-%b-%y")
    Review_Date = Review_Date.strftime("%d-%b-%y")
    Approval_Date = Approval_Date.strftime("%d-%b-%y")
except:
    logging.exception("\n \n \n Error Logged: Error in Converting Date Format of Excel File")


try:
    total_files_in_directory = os.listdir(path + "\\")
    image_file = ''
    for image in total_files_in_directory:
        if image.endswith(".jpg"):
            image_file = image
            break
        elif image.endswith(".jpeg"):
            image_file = image
            break
        elif image.endswith(".png"):
            image_file = image
            break
except:
    logging.exception("\n \n \n Error Logged: No Image file in the Folder")


try:
    total_docs = os.listdir(path + "\\Input")
    file_names = [file for file in total_docs if file.endswith('.docx')]
    file_names = [os.path.join(path + "\\Input", file) for file in file_names]
    i = 0

    for document in file_names:
        if document.endswith(".docx"):
            doc = docx.Document(document)
            ##### Code for Inserting the Image #####
            try:
                for section in doc.sections:
                    header = section.header
                    for table in header.tables:
                        table.rows[0].cells[0]._element.clear_content()
                        img = table.rows[0].cells[0].add_paragraph().add_run().add_picture(image_file, width=docx.shared.Inches(1.0), height=docx.shared.Inches(0.5))
                        table.rows[0].cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

            except:
                logging.exception("\n \n \n Error Logged: Error in Inserting the Image")


            ##### Code for Inserting Dates #####
            try:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if '<Draft Date>' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("<Draft Date>", Draft_Date)
                                elif '<Review Date>' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("<Review Date>", Review_Date)
                                elif '<Approval Date>' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("<Approval Date>", Approval_Date)

            except:
                logging.exception("\n \n \n Error Logged: Error in Inserting the Dates")



            ##### Code for Inserting the Organization Name #####
            try:
                for paragraph in doc.paragraphs:
               
                    if 'organization' in paragraph.text:
                        paragraph.text = paragraph.text.replace("<Organization Name>", Organization_Name)
            
            except:
                logging.exception("\n \n \n Error Logged: Error in Inserting the Organization Name")

            style = doc.styles['Normal']
            font = style.font
            font.name = 'Cambria'
            font.size = Pt(11)
            doc.save(output + "\\" + total_docs[i])

            i+=1
except:
    logging.exception("\n \n \n Error Logged: Error")
    pass

End_time = datetime.datetime.now()
print("Process Completed!!")
print("Time of Processing :", End_time - Start_time)
time.sleep(2)
